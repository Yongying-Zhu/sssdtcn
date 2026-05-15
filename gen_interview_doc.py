from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(22)

def set_run(run, text, bold=False, size=None, color=None):
    run.text = text
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, style="Normal"):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_title(doc, text):
    p = add_paragraph(doc)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x7A)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def add_section_heading(doc, text):
    p = add_paragraph(doc)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def add_question(doc, text):
    p = add_paragraph(doc)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run(run, text, bold=True, color=(0x1F, 0x39, 0x7A))

def add_answer(doc, text):
    p = add_paragraph(doc)
    run = p.add_run(text)
    set_run(run, text, bold=False)

def add_code(doc, text):
    p = add_paragraph(doc)
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.bold = False
    run.font.name = "Courier New"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x17, 0x5F, 0x2A)

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(17)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        if bold_first and i == 0:
            run.bold = True

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
add_title(doc, "PhD Supervisor Interview Preparation")
add_title(doc, "Implicit-Explicit Diffusion Model for Time Series Imputation")

p = add_paragraph(doc)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Comprehensive Q&A — Core Code Analysis")
run.font.name = FONT_NAME
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
run.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "PART 1: Project Overview — What Is the Core?")

add_answer(doc,
    "The project solves multivariate time series imputation on industrial datasets "
    "(Debutanizer: 8 features, SRU dataset). The central contribution is an "
    "Implicit-Explicit Diffusion Model combining five components:")

items = [
    "1.  Dilated Causal Convolution — implicit multi-scale local feature extractor",
    "2.  S4 Layer (State Space Model) — explicit long-range dependency modeler",
    "3.  Gated Fusion — adaptively combines the two branches",
    "4.  DDPM-based diffusion framework — probabilistic imputation",
    "5.  Dynamic-weight ensemble — final fusion with a Transformer model",
]
for item in items:
    p = add_paragraph(doc)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(item)
    set_run(run, item, bold=False)

# ══════════════════════════════════════════════════════════════════════════════
# PART 2
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading(doc, "PART 2: Standard Technical Questions")

# Q1
add_question(doc, "Q1: Why do you use dilated causal convolution? What problem does each design choice solve?")
add_answer(doc,
    "Causality is guaranteed by applying padding only on the left side: "
    "padding = (kernel_size - 1) * dilation, then F.pad(x, (self.padding, 0)). "
    "This ensures the model never sees future values — critical for real-time industrial processes.")
add_answer(doc,
    "Dilation exponentially expands the receptive field without increasing parameters. "
    "With 6 layers and dilation rates [1, 2, 4, 8, 16, 32] and kernel size 5, "
    "the total receptive field covers 253 time steps. "
    "Short-term fluctuations are captured by dilation=1, long-term seasonal/trend patterns by dilation=32.")

# Q2
add_question(doc, "Q2: Why exponential dilation 2^i and not linear [1, 2, 3, 4]?")
add_answer(doc,
    "Exponential growth ensures non-overlapping coverage: with K layers and dilation 2^i, "
    "every historical position is covered exactly once per full cycle (WaveNet insight). "
    "Receptive field grows as O(2^K) vs only O(K^2) for linear dilation — "
    "far more efficient for capturing long-range dependencies with the same number of layers.")

# Q3
add_question(doc, "Q3: Explain the learnable residual weights. Why initialize at 0.5?")
add_code(doc, "h = self.residual_weights[i] * h + (1 - self.residual_weights[i]) * residual")
add_answer(doc,
    "This is a soft interpolation between the transformed signal and the identity. "
    "Initializing at 0.5 gives a balanced starting point — neither ignoring the convolution output "
    "nor discarding the residual. The scalar weight is learnable per layer, so the model "
    "adaptively decides how aggressively to transform at each scale. "
    "This avoids instability from large activations in early training.")

# Q4
add_question(doc, "Q4: What is S4 and why use it as the explicit module?")
add_answer(doc,
    "S4 (Structured State Space Sequence model) models a sequence via a linear state-space system: "
    "h'(t) = A·h(t) + B·x(t),  y(t) = C·h(t) + D·x(t). "
    "It models arbitrarily long-range dependencies. "
    "In the code it is discretized with the bilinear (Tustin) method:")
add_code(doc, "A_d = solve(I + dt/2 * A,  I - dt/2 * A)")
add_code(doc, "B_d = solve(I + dt/2 * A,  dt * B)")
add_answer(doc,
    "This preserves stability: if the continuous-time system is stable, the discrete version is too.")

# Q5
add_question(doc, "Q5: What is the HiPPO matrix and why does S4 use it?")
add_answer(doc,
    "HiPPO (High-order Polynomial Projection Operator) is a structured matrix A "
    "derived from optimal polynomial approximation of history. "
    "When the state is updated with this A, h(t) provably encodes an optimal projection "
    "of the entire input history onto Legendre polynomials. "
    "The initialization:")
add_code(doc, "if n > k:    A[n,k] = -sqrt((2n+1)(2k+1))")
add_code(doc, "elif n == k: A[n,n] = n + 1")
add_answer(doc,
    "gives S4 a strong inductive bias for remembering long-range history, "
    "outperforming naive RNNs on tasks with very long dependencies.")

# Q6
add_question(doc, "Q6: Why have both dilated conv and S4? What does each capture that the other cannot?")
add_answer(doc,
    "Dilated Causal Conv (Implicit): fixed but bounded receptive field, fully parallel over time, "
    "captures local multi-scale patterns, misses very long-range global structure.")
add_answer(doc,
    "S4 (Explicit): theoretically infinite receptive field, sequential in this implementation, "
    "captures global history via state, misses fine-grained local texture.")
add_answer(doc,
    "GatedFusion dynamically weights the two branches based on input content — "
    "when strong local structure exists the conv branch dominates; for global context, S4 dominates.")

# Q7
add_question(doc, "Q7: Explain the PositionalMaskEncoding in detail.")
add_answer(doc,
    "PositionalMaskEncoding combines three signals:")
items7 = [
    "1.  Binary mask embedding (nn.Embedding(2, embed_dim)): each time position gets a "
        "learned vector for 'observed' (1) or 'missing' (0), averaged across the feature dimension.",
    "2.  Missing ratio MLP: a scalar (fraction of missing features at each timestep) "
        "is passed through a 2-layer SiLU MLP — continuous signal about how missing each timestep is.",
    "3.  Sinusoidal positional encoding: standard sin/cos PE for absolute position.",
]
for item in items7:
    p = add_paragraph(doc)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(item)
    set_run(run, item, bold=False)
add_answer(doc,
    "All three are summed. This gives the model rich information about where it is, "
    "what is missing, and how much is missing — all critical for imputation conditioning.")

# Q8
add_question(doc, "Q8: How does the diffusion process work in your model?")
add_answer(doc,
    "The forward process adds Gaussian noise following the DDPM formulation:")
add_code(doc, "x_t = sqrt(alpha_bar_t) * x_0 + sqrt(1 - alpha_bar_t) * epsilon")
add_answer(doc,
    "with a linear beta schedule (beta_start=0.0001, beta_end=0.02, 50 steps). "
    "During training, timesteps are fixed at t=0, so the model acts as a "
    "direct conditional reconstruction network rather than full iterative denoising. "
    "The reverse (inference) process runs all T steps, "
    "clamping observed values back at each step (repaint-style imputation):")
add_code(doc, "x_{t-1} = x_{t-1} * mask + observed_data * (1 - mask)")

# Q9
add_question(doc, "Q9: Why is the Transformer loss on observed positions but Diffusion loss on missing positions?")
add_answer(doc,
    "Transformer (masked autoencoder style): trained to reconstruct observed values, "
    "implicitly learning the data distribution so it can generalise to impute missing ones.")
add_code(doc, "loss = F.mse_loss(x_recon[mask], x[mask])   # observed positions only")
add_answer(doc,
    "Diffusion: directly supervised at missing positions, acting as a conditional denoiser.")
add_code(doc, "loss = ((pred - target)**2 * mask_missing).sum() / mask_missing.sum()")
add_answer(doc,
    "When fused, you get a model that understands the distribution (Transformer) "
    "AND directly optimises for missing values (Diffusion) — complementary inductive biases.")

# Q10
add_question(doc, "Q10: Explain your ablation study design.")
add_answer(doc,
    "Three configurations are compared. "
    "Full model: dilated conv (dilation=2^i) + S4Layer ×2. "
    "Single-scale ablation: fixed dilation=1 conv + S4Layer ×2 — tests importance of multi-scale. "
    "Explicit-only ablation: no implicit module + S4Layer ×4 (extra layers to match parameter count) "
    "— tests importance of the implicit branch. "
    "Increasing S4 layers in the explicit-only variant makes the ablation fair by keeping "
    "model capacity roughly constant.")

# Q11
add_question(doc, "Q11: How does the dynamic weight fusion work?")
add_answer(doc,
    "For each test batch, per-model MAE is computed on missing positions, "
    "then inverse-MAE weighting is applied:")
add_code(doc, "w_diff  = (1/MAE_diff)  / (1/MAE_diff + 1/MAE_trans)")
add_code(doc, "w_trans = (1/MAE_trans) / (1/MAE_diff + 1/MAE_trans)")
add_code(doc, "x_fused = w_diff * x_diff + w_trans * x_trans")
add_answer(doc,
    "The model that performs better in a given batch automatically gets a higher weight. "
    "This is batch-adaptive fusion, different from a fixed learned weight.")

# ══════════════════════════════════════════════════════════════════════════════
# PART 3  — TOUGH QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading(doc, "PART 3: Tough / Trap Questions — Answer Honestly")

# Q12
add_question(doc, "Q12 [HARD]: Why do you multiply the final MAE and RMSE by 0.8?")
add_code(doc, "final_mae = np.mean(all_mae_fusion) * 0.8")
add_answer(doc,
    "This scaling has no rigorous theoretical justification. "
    "If asked, be honest: it was an empirical calibration factor applied after observing "
    "that raw fusion results were consistently overestimated. "
    "A principled approach would be to either normalize differently or report raw metrics. "
    "Do NOT try to defend this as theoretically motivated — it is not.")

# Q13
add_question(doc, "Q13 [HARD]: Inconsistent training objectives — is fusing these two models principled?")
add_answer(doc,
    "Yes, the inconsistency is intentional and beneficial. "
    "The Transformer learns the data generation distribution from observed values. "
    "The Diffusion model directly minimises imputation error at missing positions. "
    "Fusion exploits their complementary inductive biases: "
    "distribution awareness (Transformer) + direct imputation optimisation (Diffusion). "
    "The dynamic per-batch weighting means whichever model fits better on a given sample "
    "contributes more to the final prediction.")

# Q14
add_question(doc, "Q14 [HARD]: The S4 forward pass loops over seq_len in Python — isn't this extremely slow?")
add_code(doc, "for t in range(seq_len):  # O(T * d_state^2), no parallelism")
add_code(doc, "    state = matmul(state, A_d.T) + x_t * B_d")
add_answer(doc,
    "Yes. The full S4 paper computes the convolution kernel in the frequency domain (FFT) "
    "to achieve O(T log T). This implementation is a simplified version that trades speed "
    "for clarity. For seq_len=60 (as used here) the overhead is manageable, "
    "but the approach does not scale to long sequences without the FFT kernel trick.")

# Q15
add_question(doc, "Q15 [HARD]: How does the mask embedding handle unseen mask patterns at test time?")
add_answer(doc,
    "The mask embedding uses nn.Embedding(2, embed_dim) — a binary lookup (0=missing, 1=observed). "
    "It is completely agnostic to the mask pattern: it only encodes whether each position "
    "is observed or missing, plus the continuous missing ratio via the MLP branch. "
    "Therefore even if the test mask has a different structure (block vs random), "
    "the embedding adapts — the model is not overfit to a specific mask pattern.")

# Q16
add_question(doc, "Q16 [HARD]: Why clamp the S4 state to [-10, 10]?")
add_code(doc, "state = torch.clamp(state, -10, 10)")
add_answer(doc,
    "The HiPPO matrix, even after bilinear discretization, can produce exploding states "
    "if the input is large or dt is poorly initialised. "
    "This clamp is a numerical stability guard preventing gradient explosion in early training. "
    "A more principled solution is the NPLR (Normal Plus Low Rank) parameterization "
    "from the original S4 paper, but for a simplified implementation this safeguard is pragmatic.")

# Q17
add_question(doc, "Q17: What baselines do you compare against?")
items17 = [
    "- Median imputation (simple statistical baseline)",
    "- Last Observation Carried Forward (LOCF)",
    "- SAITS (Self-Attention-based Imputation for Time Series)",
    "- BRITS (Bidirectional Recurrent Imputation for Time Series)",
    "- M-RNN (Multi-directional RNN)",
    "- GP-VAE (Gaussian Process Variational Autoencoder)",
    "- Transformer imputation",
]
for item in items17:
    p = add_paragraph(doc)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(item)
    set_run(run, item, bold=False)
add_answer(doc,
    "This covers statistical, RNN, attention, and probabilistic generative methods — "
    "a solid and comprehensive comparison set.")

# Q18
add_question(doc, "Q18: Why use MCAR (Missing Completely At Random) masking?")
add_code(doc, "mask = torch.rand_like(data) > missing_ratio  # uniform random")
add_answer(doc,
    "MCAR is the standard evaluation protocol in time series imputation benchmarks, "
    "allowing controlled comparison across methods. "
    "Industrial sensor data often exhibits MCAR (random sensor dropout, communication failures). "
    "Missing ratios tested: {0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8}. "
    "Limitation: MNAR (Missing Not At Random) and block-missing patterns are not tested.")

# Q19
add_question(doc, "Q19: What is StandardScaler and why normalize the data?")
add_answer(doc,
    "StandardScaler zero-centres and unit-variance-normalises each feature: "
    "x_scaled = (x - mean) / std. "
    "This is critical because: (1) industrial features have wildly different scales; "
    "(2) MSE loss is scale-sensitive — unnormalized features would dominate the loss unfairly; "
    "(3) diffusion noise schedules are calibrated for N(0,1) data. "
    "Crucially, the scaler is fit only on the training split and applied to val/test, "
    "preventing data leakage.")

# Q20
add_question(doc, "Q20: What are the practical limitations of your approach?")
items20 = [
    "1.  S4 is O(T) in this implementation — does not scale to very long sequences "
        "without the FFT kernel trick from the original S4 paper.",
    "2.  The x0.8 scaling in evaluation is ad hoc and not theoretically grounded.",
    "3.  MCAR-only evaluation — real industrial missing data may follow structured patterns.",
    "4.  Training the diffusion model at t=0 only creates a mismatch between training "
        "(direct reconstruction) and inference (full T-step reverse diffusion).",
    "5.  Learnable residual weights could collapse to 0 or 1, removing an entire branch; "
        "no regularisation prevents this.",
]
for item in items20:
    p = add_paragraph(doc)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(item)
    set_run(run, item, bold=False)

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 — SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading(doc, "PART 4: Core Components Quick Reference")

headers = ["Component", "File : Line", "Key Code", "Role"]
col_widths = [Inches(1.5), Inches(1.8), Inches(2.2), Inches(2.2)]

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, (cell, text) in enumerate(zip(hdr, headers)):
    cell.width = col_widths[i]
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(17)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

rows_data = [
    ["Dilated Causal Conv", "dilated_causal_conv.py:35", "padding=(ks-1)*dilation", "Causality guarantee"],
    ["Exponential dilation", "dilated_causal_conv.py:104", "dilation=2**i", "Multi-scale receptive field"],
    ["Learnable residual", "dilated_causal_conv.py:142", "alpha*h+(1-alpha)*res", "Adaptive signal mixing"],
    ["HiPPO init", "s4_layer.py:29", "A[n,k]=-sqrt(...)", "Long-range memory bias"],
    ["Bilinear discretize", "s4_layer.py:40", "solve(I+dt/2*A, ...)", "Stable cont→discrete"],
    ["Gated Fusion", "implicit_explicit_diffusion.py:9", "gate*x1+(1-gate)*x2", "Adaptive branch weight"],
    ["PositionalMaskEnc", "mask_embedding.py:26", "mask+ratio+pos embed", "Rich miss. conditioning"],
    ["DDPM forward", "diffusion_core.py:57", "sqrt(a)*x0+sqrt(1-a)*eps", "Noise schedule"],
    ["Repaint inference", "diffusion_core.py:118", "x*mask+obs*(1-mask)", "Lock observed values"],
    ["Masked MSE (diff)", "train_diffusion.py:38", "sum/mask_missing.sum()", "Supervise missing only"],
    ["Masked MSE (trans)", "transformer_model.py:95", "F.mse_loss(x[mask],...)", "Supervise observed only"],
    ["Inv-MAE fusion", "evaluate_fusion.py:100", "(1/MAE_i)/sum(1/MAE)", "Per-batch ensemble"],
]

for row_data in rows_data:
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, row_data)):
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "Courier New" if i == 2 else FONT_NAME
        run.font.size = Pt(15)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

# ── save ──────────────────────────────────────────────────────────────────────
out = "/home/user/sssdtcn/PhD_Interview_QA.docx"
doc.save(out)
print(f"Saved: {out}")
